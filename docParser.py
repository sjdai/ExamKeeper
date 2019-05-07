import docx

questionType = ['vocabulary','multiplechoice','cloze','passagecompletion','readingcomprehension','fitthebestsentence','translation']
optionType = ['(A)','(B)','(C)','(D)','(E)','(F)','(G)','(H)','(I)','(J)','(K)','(L)','(M)','(N)','(O)','(P)','(Q)','(R)','(S)','(T)','(U)','(V)','(W)','(X)','(Y)','(Z)']

def retrieveIndex(context):
    idx = []
    if 'are based on the following passage.' in context:
        context = context.replace('-',' ')
        context = context.split()
        idx = [i for i in context if i.isdigit()]
    return idx

def isOption(context):
    val = False
    for o in optionType:
        if o in context:
            val = True
    return val

def indexer(context,point):
    paragraphs = [i.replace(' ','') for i in context]
    index = [[i,paragraphs.index(i)]for i in point if i in paragraphs]
    return index

def ParseVoc(context,name):
    context = [i for i in context if i != '']
    vocAnswer = ans['Vocabulary']
    questions = []
    context.pop(0)

    for term in range(len(context)):
        tmp = context[term]
        tmp = tmp.split()
        for t in tmp:
            if '_' in t:
                tmp[tmp.index(t)] = vocAnswer[term]
        questions.append({'sentence':' '.join(tmp),'answer':vocAnswer[term]})

    return questions

def ParseMultiple(context,name):
    context = [i for i in context if i != '']
    Answer = ans[name]
    answerId = 0
    questions = []
    context.pop(0)
    article = ''
    for term in range(len(context)):
        tmp = context[term]
        tmp = tmp.split()
        key = isOption(context[term])
        if key:
            index = indexer(tmp,optionType)
            for i in index:
                tmp[i[1]] = ''
            tmp = [i for i in tmp if i != '']
            tmp.pop(0)
            questions.append({'article':article,'answer':Answer[answerId],'option':tmp})
            answerId += 1
        else:
            for t in tmp:
                if '_' in t:
                    tmp[tmp.index(t)] = '_'
                    article = ' '.join(tmp)
    return questions

def ParseCloze(context,name):
    context = [i for i in context if i != '']
    Answer = ans[name]
    answerId = 0
    questions = []
    context.pop(0)
    article = ''
    context.pop(0)
    option = []
    answer = []
    appendFlag = False
    for term in range(len(context)):
        tmp = context[term]
        tmp = tmp.split()
        key = isOption(context[term])
        if key:
            appendFlag = True
            index = indexer(tmp,optionType)
            for i in index:
                tmp[i[1]] = ''
            tmp = [i for i in tmp if i != '']
            tmp.pop(0)
            option.append(tmp)
        else:
            if appendFlag:
                appendFlag = False
                questions.append({'article':article,'answer':answer,'option':option})
                article = ''
                option = []
                answer = []
            for t in tmp:
                if '_' in t:
                    tmp[tmp.index(t)] = '_'
                    answer.append(Answer[answerId])
                    answerId += 1
            article += ' '.join(tmp)
    if appendFlag == True or len(questions) == 0:
        questions.append({'article':article,'answer':answer,'option':option})

    return questions

def ParsePassageCompletion(context,name):
    context = [i for i in context if i != '']
    Answer = ans[name]
    answerId = 0
    questions = []
    context.pop(0)
    article = ''
    context.pop(0)
    option = []
    answer = []
    appendFlag = False
    for term in range(len(context)):
        tmp = context[term]
        tmp = tmp.split()
        key = isOption(context[term])
        if key:
            appendFlag = True
            index = indexer(tmp,optionType)
            for i in index:
                tmp[i[1]] = ''
            tmp = [i for i in tmp if i != '']
            option.extend(tmp)
        else:
            if appendFlag:
                appendFlag = False
                questions.append({'article':article,'answer':answer,'option':option})
                article = ''
                option = []
                answer = []
            for t in tmp:
                if '_' in t:
                    tmp[tmp.index(t)] = '_'
                    answer.append(Answer[answerId])
                    answerId += 1
            article += ' '.join(tmp)
    if appendFlag == True or len(questions) == 0:
        questions.append({'article':article,'answer':answer,'option':option})

    return questions

def ParseFt(context,name):
    context = [i for i in context if i != '']
    Answer = ans[name]
    questions = []
    context.pop(0)
    answerId = 0
    appendFlag = False
    article = ''
    option = []
    questionId = 0
    label = 0
    for term in range(len(context)):
        tmp = context[term]
        key = isOption(tmp)
        tmp = tmp.split()
        if key:
            appendFlag = True
            index = indexer(tmp,optionType)
            for i in index:
                tmp[i[1]] = ''
            tmp = [i for i in tmp if i != '']
            tmp = ' '.join(tmp)
            option.append(tmp)
        else:
            if appendFlag:
                appendFlag = False
                article = article.split()
                questionId = article.pop(0)
                article = ' '.join(article)
                label = ord(Answer[answerId]) - 65
                questions.append({'id':questionId,'context_sentence':article,'start_ending':'','ending_0': option[0],'ending_1': option[1],'ending_2': option[2],'ending_3': option[3],'label': label})
                option = []
                article = ''
                answerId += 1

            for t in tmp:
                if '_' in t:
                    tmp[tmp.index(t)] = '_'
            article += ' '.join(tmp)
    if appendFlag == True or len(questions) == 0:
        article = article.split()
        questionId = article.pop(0)
        article = ' '.join(article)
        label = ord(Answer[answerId]) - 65

        questions.append({'id':questionId,'context_sentence':article,'start_ending':'','ending_0': option[0],'ending_1': option[1],'ending_2': option[2],'ending_3': option[3],'label': label})

    return questions

def ParseRc(context,name):
    context = [i for i in context if i != '']
    Answer = ans[name]
    answerId = 0
    question = []
    questions = []
    context.pop(0)
    article = ''
    options = []
    option = []
    appendFlag = False
    start = '1'
    questionId = []
    for term in range(len(context)):
        tmp = context[term]
        reIdx = retrieveIndex(tmp)
        key = isOption(tmp)
        if len(reIdx) > 0:
            start = reIdx[0]
            end = reIdx[1]
            continue
        if start+'.' in tmp:
            tmp = tmp.replace(start+'.','')
            question.append(tmp)
            questionId.append(start)
            start = str(int(start)+1)

        elif key:
            appendFlag = True
            tmp = tmp.split()
            index = indexer(tmp,optionType)
            index.append(['',len(tmp)])
            for i in range(1,len(index)):
                array = tmp[index[i-1][1]:index[i][1]]
                array = [i for i in array if i != '']
                array.pop(0)
                array = ' '.join(array)
                option.append(array)
        else:
            tmp = tmp.split()
            if appendFlag:
                appendFlag = False
                optionNum = int(len(option)/len(question))
                op = 0
                for j in range(1,len(question)+1):
                    questions.append({'id':questionId[j-1],'article':article,'truth':ord(Answer[answerId])-65,'question':question[j-1],'option':option[op:op+optionNum]})
                    answerId += 1
                    op += optionNum
                op = 0
                article = ''
                option = []
                question = []
            article += ' '.join(tmp)
    if appendFlag == True or len(questions) == 0:
        optionNum = int(len(option)/len(question))
        op = 0
        for j in range(len(question)):
            questions.append({'id':questionId[j],'article':article,'truth':ord(Answer[answerId])-65,'question':question[j],'option':option[op:op+optionNum]})
            answerId +=1
            op += optionNum

    return questions

def ParseTr(context,name):
    context = [i for i in context if i != '']
    Answer = ans[name]
    questions = []
    context.pop(0)
    idx = 0
    for term in context:
        if 'This could be any language other than English' in term:
            continue
        questions.append({'sentence':term,'answer':Answer[idx]})
        idx+=1
    return questions

def main(filename):
    doc = docx.Document(filename)
    paragraphs = [para.text.lower() for para in doc.paragraphs]
    paragraphs = [i.replace(' ','') for i in paragraphs]
    index = [[i,paragraphs.index(i)]for i in questionType if i in paragraphs]
    paragraphs = [para.text for para in doc.paragraphs]
    index.append(['',len(paragraphs)])
    for idx in range(1,len(index)):
        tp = index[idx-1][0]
        qetype = paragraphs[index[idx-1][1]]
        qetype = ' '.join(qetype.split())
        parseArray = paragraphs[index[idx-1][1]:index[idx][1]]
        if tp == 'vocabulary':
            print(ParseVoc(parseArray,qetype))
        elif tp == 'multiplechoice':
            print(ParseMultiple(parseArray,qetype))
        elif tp == 'cloze':
            print(ParseCloze(parseArray,qetype))
        elif tp == 'passagecompletion':
            print(ParsePassageCompletion(parseArray,qetype))
        elif tp == 'readingcomprehension':
            print(ParseRc(parseArray,qetype))
        elif tp == 'translation':
            print(ParseTr(parseArray,qetype))
        elif tp == 'fitthebestsentence':
            print(ParseFt(parseArray,qetype))
        elif tp == '':
            pass
        else:
            return 'error occured!'

def answer(filename):
    ans = dict()
    doc = docx.Document(filename)

    paragraphs = [para.text.lower() for para in doc.paragraphs]
    paragraphs = [i.replace(' ','') for i in paragraphs]
    index = [[i,paragraphs.index(i)]for i in questionType if i in paragraphs]
    paragraphs = [para.text for para in doc.paragraphs]
    index.append(['',len(paragraphs)])
    for idx in range(1,len(index)):
        tp = index[idx-1][0]
        qetype = paragraphs[index[idx-1][1]]
        parseArray = paragraphs[index[idx-1][1]:index[idx][1]]
        parseArray.pop(0)
        ans[qetype] = [pa.split('.')[0] for pa in parseArray]
    return ans
ans = answer('answerKey.docx')
print(ans)
main('exam.docx')

